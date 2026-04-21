"""
팀원 설명용 DOCX 생성 스크립트
실행: python make_team_doc.py
출력: 팩트체크_카테고리분류_팀원설명.docx
"""

import sys
sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 페이지 여백 설정 ─────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(3.0)

# ── 스타일 헬퍼 ─────────────────────────────────────────────
def set_font(run, size=11, bold=False, color=None, name="맑은 고딕"):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), name)
    rPr.insert(0, rFonts)

def add_heading(doc, text, level=1, color=(0, 70, 127)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    sizes = {1: 16, 2: 13, 3: 11}
    set_font(run, size=sizes.get(level, 11), bold=True, color=color)
    return p

def add_body(doc, text, indent=0, bullet=False, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    prefix = "• " if bullet else ""
    run = p.add_run(prefix + text)
    set_font(run, size=10)
    return p

def add_bold_body(doc, label, text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    r1 = p.add_run(label)
    set_font(r1, size=10, bold=True, color=(0, 70, 127))
    r2 = p.add_run(text)
    set_font(r2, size=10)
    return p

def shade_cell(cell, hex_color="DEEAF1"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, size=10, color=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)

def add_table(doc, headers, rows, col_widths=None, header_color="1F4E79"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        shade_cell(cell, header_color.replace("#", ""))
        set_cell_text(cell, h, bold=True, size=10,
                      color=(255,255,255), align="center")

    # 데이터
    for i, row in enumerate(rows):
        bg = "F2F2F2" if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(row):
            cell = table.rows[i+1].cells[j]
            shade_cell(cell, bg)
            set_cell_text(cell, val, size=9)

    if col_widths:
        for row in table.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = Cm(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run("─" * 55)
    set_font(run, size=9, color=(180, 180, 180))

def add_code_block(doc, lines):
    """회색 배경 코드 블록 (단락 여러 개)."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.5)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.space_before = Pt(1)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "F0F0F0")
        pPr.append(shd)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)


# ══════════════════════════════════════════════════════════════
# 표지
# ══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run("삼선뉴스 (三鮮)")
set_font(run, size=22, bold=True, color=(0, 70, 127))

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("팩트체크 파이프라인 & 카테고리 자동 분류\n팀원 기술 공유 문서")
set_font(run2, size=13, color=(80, 80, 80))

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(8)
run3 = p3.add_run("담당: 이동우 · 2026-04-15")
set_font(run3, size=10, color=(130, 130, 130))

doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# 1. 문서 목적
# ══════════════════════════════════════════════════════════════
add_heading(doc, "1. 이 문서의 목적", level=1)
add_body(doc, "이번 스프린트(5주차)에서 구현한 두 가지 기능을 팀원들이 빠르게 이해하고 각자 파트에 연동할 수 있도록 정리합니다.")
add_body(doc, "팩트체크 파이프라인 (Step 0~2 완성, Step 3 구현 예정)", indent=0.8, bullet=True)
add_body(doc, "기사 카테고리 자동 분류기 (키워드 기반, 표준 카테고리 11종)", indent=0.8, bullet=True)


# ══════════════════════════════════════════════════════════════
# 2. 팩트체크 파이프라인
# ══════════════════════════════════════════════════════════════
add_heading(doc, "2. 팩트체크 파이프라인", level=1)

add_heading(doc, "2-1. 왜 만드나요?", level=2)
add_body(doc, "기획서 8-4절의 논문 기반 설계를 코드로 구현한 것입니다. 핵심 원칙 세 가지:")
add_body(doc, "LLM 단독 판단 금지 — Bad Actor Good Advisor(AAAI 2024): LLM이 직접 팩트 여부를 판단하면 F1 0.57~0.72에 불과합니다. LLM은 판단자(Judge)가 아닌 분석 어드바이저(Advisor)로만 활용합니다.", indent=0.8, bullet=True)
add_body(doc, "3단계 라벨 — LIAR(ACL 2017): True/False 이진 분류 대신 FACT / RUMOR / UNVERIFIED 3단계로 압축합니다. 현실 뉴스에는 회색지대가 많기 때문입니다.", indent=0.8, bullet=True)
add_body(doc, "증거 기반 판단 — COLING 2025: 외부 증거 없이 LLM이 혼자 판단하면 오분류율이 높습니다. Google Search Grounding으로 실시간 증거를 주입합니다.", indent=0.8, bullet=True)

add_heading(doc, "2-2. 전체 처리 흐름", level=2)
add_body(doc, "기사 1건이 DB에 저장되기 전에 아래 순서로 자동 처리됩니다.")

steps = [
    ["Step", "이름", "방식", "비용/속도", "결과"],
    ["0", "채널 등급 분류", "소스별 정적 룩업 + 의견/사설 신호 탐지", "규칙 기반, 무료", "DROP 또는 계속"],
    ["1", "루머·유출 신호 탐지", "한국어+영어 정규식 패턴 매칭", "규칙 기반, 무료", "FACT_AUTO / RUMOR / 계속"],
    ["2", "Google FC API", "ClaimReview DB 200개+ 기관 조회", "무료, ~200ms", "라벨 확정 or UNVERIFIED"],
    ["3", "Gemini 2-Pass", "Pass A(문체) + Pass B(상식) + Search Grounding", "~$0.001/건, 구현 예정", "FACT / RUMOR / UNVERIFIED"],
]
add_table(doc, steps[0], steps[1:], col_widths=[1.2, 3.5, 5.0, 3.5, 3.5])

add_heading(doc, "2-3. 채널 등급 5종 (Step 0)", level=2)
add_body(doc, "11개 채널을 5개 등급으로 사전 분류합니다. 등급에 따라 DB 저장 여부와 팩트 처리 방식이 달라집니다.")

tiers = [
    ["등급", "대상 채널 예시", "처리 방식", "앱 노출"],
    ["언론사 🟢 Official", "MIT TR, IEEE, TechCrunch 등", "루머 신호 없으면 FACT 자동 확정", "메인 피드 기본 노출"],
    ["언론사 🟡 Credible Leak", "VentureBeat AI (기본값)", "자동 RUMOR + Google FC → Gemini 검증", "[루머/유출] 뱃지"],
    ["언론사 🔴 Opinion", "사설·칼럼 감지 시", "즉시 DROP — DB 진입 차단", "노출 안 함"],
    ["커뮤니티 🟢 High-Signal", "Reddit r/ML, arXiv·GitHub 링크 포함", "Gemini 검증 후 최종 라벨", "[AI 트렌드] 뱃지"],
    ["커뮤니티 🔴 Noise", "밈·질문·출처 없는 토론", "즉시 DROP — DB 진입 차단", "노출 안 함"],
]
add_table(doc, tiers[0], tiers[1:], col_widths=[4.0, 4.5, 5.5, 3.5])

add_heading(doc, "2-4. 루머 신호 사전 (Step 1)", level=2)
add_body(doc, "제목 + 본문 앞 1,000자를 스캔합니다. 한국어·영어 모두 지원합니다.")

signals = [
    ["신호 강도", "영어 패턴 예시", "한국어 패턴 예시", "처리"],
    ["강한 신호 → RUMOR", "allegedly, reportedly, unverified, purportedly", "루머, 소문, ~라는 주장, 허위, 가짜 뉴스", "즉시 RUMOR 라벨"],
    ["약한 신호 → NEEDS_VERIFICATION", "sources say, expected to, may be, possibly", "소식통에 따르면, 단독, ~로 보인다, 전해졌다", "FC API → Gemini"],
    ["의견 신호 → DROP", "opinion, editorial, I think, I believe", "칼럼, 사설, 기고, 내 생각", "즉시 DROP"],
]
add_table(doc, signals[0], signals[1:], col_widths=[4.0, 5.5, 5.5, 3.0])

add_heading(doc, "2-5. Google FC API 활성화 방법 (Step 2)", level=2)
add_body(doc, "Google Fact Check Tools API는 완전 무료입니다. API 키만 발급하면 바로 동작합니다.")
add_bold_body(doc, "① ", "Google Cloud Console → 새 프로젝트 생성", indent=0.8)
add_bold_body(doc, "② ", "'Fact Check Tools API' 검색 → 사용 설정", indent=0.8)
add_bold_body(doc, "③ ", "사용자 인증 정보 → API 키 발급", indent=0.8)
add_bold_body(doc, "④ ", ".env 파일에 추가:", indent=0.8)
add_code_block(doc, ["GOOGLE_FC_API_KEY=발급받은_키_여기에_붙여넣기"])
add_body(doc, "API 키가 없으면 Step 2를 자동으로 건너뛰고 UNVERIFIED로 처리합니다. (파이프라인 크래시 없음)")


# ══════════════════════════════════════════════════════════════
# 3. 팀원별 연동 포인트
# ══════════════════════════════════════════════════════════════
add_heading(doc, "3. 팀원별 연동 포인트", level=1)

members = [
    ["팀원", "역할", "팩트체크 연동 포인트", "필요한 액션"],
    ["이상준", "RSS 크롤러", "source, source_type 필드 정확히 설정하면 Step 0 자동 동작", "기존 코드 변경 없음 ✅"],
    ["강주찬", "백엔드 / Supabase", "save_articles() 호출 시 파이프라인 자동 실행됨", "fact_checks 테이블 스키마 확인"],
    ["정수민", "프론트엔드", "articles.fact_label 값으로 뱃지 렌더링", "FACT / RUMOR / UNVERIFIED 뱃지 UI 구현"],
    ["김민규", "PM", "Google FC API 키 발급 필요", "Google Cloud Console에서 발급 후 .env 전달"],
    ["이동우", "AI 모델", "Step 3 Gemini 2-Pass 구현 예정", "팀 확인 후 진행"],
]
add_table(doc, members[0], members[1:], col_widths=[2.0, 3.5, 6.0, 4.0])

add_heading(doc, "3-1. DB 저장 흐름 (강주찬님 참고)", level=2)
add_body(doc, "save_articles.py가 이미 업데이트되어 있습니다. 별도 코드 수정 없이 자동으로 아래 흐름으로 동작합니다.")
add_code_block(doc, [
    "save_articles(articles) 호출",
    "  └─ for each article:",
    "       ├─ run_fact_check(title, content, source, source_type)",
    "       │    └─ Step 0 → Step 1 → Step 2",
    "       ├─ DROP이면 DB 저장 스킵",
    "       ├─ articles.fact_label = 결과 라벨",
    "       └─ save_fact_checks(url_hash, claims)  ← fact_checks 테이블",
])

add_heading(doc, "3-2. 프론트엔드 뱃지 (정수민님 참고)", level=2)
add_body(doc, "articles.fact_label 컬럼 값에 따라 카드 뱃지를 다르게 표시해 주세요.")

badges = [
    ["fact_label 값", "뱃지 텍스트", "색상 가이드", "설명"],
    ["FACT", "(뱃지 없음)", "기본", "공신력 언론사 검증 완료 기사"],
    ["RUMOR", "🔴 루머/유출", "빨간색", "루머 신호 감지 또는 유출 보도"],
    ["UNVERIFIED", "❓ 미검증", "노란색", "아직 검증 중 (Gemini 대기 포함)"],
    ["DROP", "표시 안 함", "-", "의견/노이즈 — DB에 저장 안 됨"],
]
add_table(doc, badges[0], badges[1:], col_widths=[3.5, 3.5, 3.0, 6.0])


# ══════════════════════════════════════════════════════════════
# 4. 카테고리 자동 분류기
# ══════════════════════════════════════════════════════════════
add_heading(doc, "4. 카테고리 자동 분류기", level=1)

add_heading(doc, "4-1. 배경", level=2)
add_body(doc, "기존에는 RSS 피드 소스 단위로 카테고리가 고정되어 있었습니다 (예: VentureBeat AI의 모든 기사 → 'AI 비즈니스').")
add_body(doc, "기사 내용을 보지 않고 소스 이름으로만 분류 → 실제 내용과 다른 카테고리가 붙는 문제", indent=0.8, bullet=True)
add_body(doc, "MIT Tech Review(ai_only=False)처럼 AI 무관 기사도 섞이는 피드에서 더 심각", indent=0.8, bullet=True)
add_body(doc, "이제 기사별로 제목 + 본문 내용을 읽고 카테고리를 자동으로 결정합니다.")

add_heading(doc, "4-2. 동작 방식", level=2)
add_bold_body(doc, "1단계 키워드 매칭 ", "(무료, 즉시): 11개 카테고리별 키워드 사전으로 점수를 계산합니다. 제목에 3배 가중치를 적용합니다.")
add_bold_body(doc, "2단계 폴백 ", "(점수 4점 미만): 소스 기본 카테고리를 그대로 유지합니다.")
add_bold_body(doc, "Ollama 폴백 ", "(선택): use_ollama=True 옵션 시 qwen3.5:4b 모델로 분류합니다. 기본 비활성.")

add_heading(doc, "4-3. 표준 카테고리 11종 (앱 피드 필터 기준)", level=2)
add_body(doc, "프론트엔드 피드 필터 탭은 아래 카테고리 기준으로 구현해 주세요.")

cats = [
    ["카테고리", "주요 키워드 예시", "해당 소스 예시"],
    ["LLM·언어모델", "GPT, Claude, Llama, RLHF, fine-tuning", "TechCrunch, The Decoder"],
    ["AI 반도체·하드웨어", "GPU, chip, Nvidia, H100, NPU, TSMC", "IEEE Spectrum, The Verge"],
    ["AI 모델·연구", "model release, benchmark, multimodal, diffusion", "MIT Tech Review"],
    ["AI 비즈니스·투자", "funding, startup, series A, IPO, acquisition", "VentureBeat AI, TechCrunch"],
    ["AI 윤리·정책", "regulation, EU AI Act, bias, copyright, ethics", "The Guardian Tech"],
    ["AI 서비스·제품", "launch, app, chatbot, API, feature", "Product Hunt, TechCrunch"],
    ["AI 오픈소스", "open source, GitHub, Hugging Face, MIT license", "Reddit r/LocalLLaMA"],
    ["AI 에이전트·자동화", "agent, autonomous, tool use, workflow", "The Decoder, VentureBeat AI"],
    ["AI 로보틱스", "robot, humanoid, autonomous vehicle, drone", "IEEE Spectrum"],
    ["AI 연구·논문", "arXiv, paper, NeurIPS, ICML, benchmark", "Reddit r/MachineLearning"],
    ["테크 일반", "smartphone, social media, gaming, blockchain", "The Verge, MIT Tech Review"],
]
add_table(doc, cats[0], cats[1:], col_widths=[4.0, 6.5, 5.5])

add_heading(doc, "4-4. 분류 결과 예시", level=2)

examples = [
    ["입력 제목 (영문)", "소스 기본값", "분류 결과"],
    ["Nvidia launches H200 GPU with 141GB HBM3e", "AI/반도체", "AI 반도체·하드웨어"],
    ["OpenAI releases GPT-5 with 1M token context", "AI 심층", "LLM·언어모델"],
    ["Meta open-sources Llama 4 on GitHub (MIT)", "AI 비즈니스", "AI 오픈소스"],
    ["EU AI Act regulation: what companies must do", "AI 윤리", "AI 윤리·정책"],
    ["Humanoid robot walks 10km without falling", "테크 전반", "AI 로보틱스"],
    ["arXiv: GPT-4 struggles with commonsense reasoning", "AI 연구", "AI 연구·논문"],
]
add_table(doc, examples[0], examples[1:], col_widths=[6.5, 3.5, 4.5])


# ══════════════════════════════════════════════════════════════
# 5. 파일 구조 요약
# ══════════════════════════════════════════════════════════════
add_heading(doc, "5. 추가된 파일 구조", level=1)

add_code_block(doc, [
    "samseon/",
    "├── collect/",
    "│   └── classifier.py              ← 카테고리 자동 분류기 (신규)",
    "│       └── classify(title, content, fallback) → 카테고리 문자열",
    "│",
    "├── fact_checker/                  ← 팩트체크 패키지 (신규)",
    "│   ├── __init__.py",
    "│   ├── channel_config.py          ← 11개 채널 등급 + credibility_score 정의",
    "│   ├── signal_detector.py         ← 한/영 루머·의견·유출 신호 탐지 (Step 0~1)",
    "│   ├── google_fc_api.py           ← Google FC API 래퍼 (Step 2)",
    "│   └── pipeline.py                ← Step 0~2 오케스트레이터",
    "│       └── run_fact_check(title, content, source, source_type) → FactCheckResult",
    "│",
    "└── backend/",
    "    └── save_articles.py           ← 팩트체크 파이프라인 자동 연결 (수정됨)",
])

add_divider(doc)
add_heading(doc, "6. 현재 구현 현황 & 다음 단계", level=1)

status = [
    ["항목", "상태", "비고"],
    ["채널 등급 분류 (Step 0)", "✅ 완료", "channel_config.py"],
    ["루머·의견 신호 탐지 (Step 1)", "✅ 완료", "한국어+영어 패턴 사전 포함"],
    ["Google FC API (Step 2)", "✅ 완료", ".env에 GOOGLE_FC_API_KEY 추가 시 활성화"],
    ["Gemini 2-Pass Advisor (Step 3)", "⏳ 구현 예정", "팀 확인 후 이동우가 진행"],
    ["save_articles.py 연결", "✅ 완료", "별도 호출 코드 불필요"],
    ["카테고리 자동 분류기", "✅ 완료", "collect/classifier.py"],
    ["rss_crawler.py 연결", "✅ 완료", "기사별 분류 자동 적용"],
]
add_table(doc, status[0], status[1:], col_widths=[5.5, 2.5, 8.0])

p_end = doc.add_paragraph()
p_end.paragraph_format.space_before = Pt(12)
r_end = p_end.add_run("문의: 이동우 (Discord / KakaoTalk)")
set_font(r_end, size=9, color=(130, 130, 130))

# ── 저장 ────────────────────────────────────────────────────
out_path = r"C:\Users\이동우\samseon\팩트체크_카테고리분류_팀원설명.docx"
doc.save(out_path)
print(f"저장 완료: {out_path}")
