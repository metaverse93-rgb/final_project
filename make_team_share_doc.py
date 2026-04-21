"""
make_team_share_doc.py — 담당별 공유 사항 정리 문서
실행: python make_team_share_doc.py
출력: 팀원_공유사항_담당별정리.docx
"""

import sys
sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(3.0)


# ── 스타일 헬퍼 ─────────────────────────────────────────────
def set_font(run, size=11, bold=False, color=None, name="맑은 고딕"):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), name)
    rPr.insert(0, rFonts)

def add_heading(doc, text, level=1, color=(0, 70, 127)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    sizes = {1: 16, 2: 13, 3: 11}
    set_font(run, size=sizes.get(level, 11), bold=True, color=color)
    return p

def add_body(doc, text, indent=0, bullet=False, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    prefix = "• " if bullet else ""
    run = p.add_run(prefix + text)
    set_font(run, size=10)
    return p

def shade_cell(cell, hex_color="DEEAF1"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, size=10, color=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)

def add_table(doc, headers, rows, col_widths=None, header_color="1F4E79"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        shade_cell(cell, header_color)
        set_cell_text(cell, h, bold=True, size=10, color=(255,255,255), align="center")
    for i, row in enumerate(rows):
        bg = "F2F2F2" if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(row):
            cell = table.rows[i+1].cells[j]
            shade_cell(cell, bg)
            set_cell_text(cell, val, size=9)
    if col_widths:
        for row in table.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table

def add_code_block(doc, lines, bg="F0F0F0"):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.5)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.space_before = Pt(1)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  bg)
        pPr.append(shd)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)

def add_callout(doc, text, bg="FFF2CC", color=(127, 96, 0)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  bg)
    pPr.append(shd)
    run = p.add_run(text)
    set_font(run, size=10, color=color)
    return p

def add_person_heading(doc, name, role, color):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  color)
    pPr.append(shd)
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run(f"  {name}  |  {role}")
    set_font(run, size=13, bold=True, color=(255, 255, 255))
    return p

def add_action_label(doc, text, urgent=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    color = (180, 0, 0) if urgent else (100, 100, 0)
    bg    = "FDECEA"   if urgent else "FFFDE7"
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  bg)
    pPr.append(shd)
    run = p.add_run(f"  {text}")
    set_font(run, size=10, bold=True, color=color)
    return p

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run("─" * 55)
    set_font(run, size=9, color=(180, 180, 180))


# ══════════════════════════════════════════════════════════════
# 표지
# ══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run("삼선뉴스 (三鮮)")
set_font(run, size=22, bold=True, color=(0, 70, 127))

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("팀원 공유 사항 — 담당별 정리")
set_font(run2, size=14, color=(80, 80, 80))

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(8)
run3 = p3.add_run("이동우 작업분 기준  ·  2026-04-21  |  6주차")
set_font(run3, size=10, color=(130, 130, 130))

doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# 전체 변경 요약 (1페이지)
# ══════════════════════════════════════════════════════════════
add_heading(doc, "전체 변경 요약", level=1)
add_body(doc, "이번 작업에서 변경/추가된 내용을 팀원 담당별로 정리합니다. 각자 해당 섹션을 확인하고 액션이 필요한 항목을 처리해주세요.")

add_table(doc,
    ["담당", "액션 필요", "참고 사항"],
    [
        ["강주찬", "DB 컬럼 추가 (ALTER TABLE)\nOllama 모델 교체", "임베딩 모델 통일\n어드민 API 신규 추가"],
        ["정수민", "UNVERIFIED 기사 UI 처리 방식\n(PM 협의 후 결정)", "needs_review 필드 API 응답에 포함됨\n어드민 API 연동 가능"],
        ["이상준", "source 필드값 channel_config 이름과 일치 확인", "classifier.py 신규 추가\ntranslator/summarizer.py 삭제"],
        ["김민규 (PM)", "UNVERIFIED 기사 피드 노출 정책 결정\nVentureBeat API 비용 증가 인지", "전체 변경 내역 파악"],
    ],
    col_widths=[3.0, 5.5, 6.0]
)


# ══════════════════════════════════════════════════════════════
# 강주찬
# ══════════════════════════════════════════════════════════════
add_person_heading(doc, "강주찬", "RAG + 백엔드 + Supabase", "1F4E79")

add_action_label(doc, "🔴 액션 필요 ①  —  Supabase 스키마 변경", urgent=True)
add_body(doc, "Supabase > SQL Editor에서 아래 SQL을 실행해주세요.")
add_code_block(doc, [
    "ALTER TABLE articles ADD COLUMN needs_review BOOLEAN DEFAULT FALSE;",
])
add_body(doc, "추가 이유: UNVERIFIED 판정 기사를 관리자가 수동 검토할 수 있도록 대기 플래그 컬럼이 필요합니다.", indent=0.5)

add_action_label(doc, "🔴 액션 필요 ②  —  Ollama 임베딩 모델 교체", urgent=True)
add_body(doc, "임베딩 모델이 아래와 같이 전체 통일됐습니다. 로컬에서 새 모델을 pull해주세요.")
add_code_block(doc, [
    "ollama pull qwen3-embedding:0.6b",
])
add_table(doc,
    ["파일", "변경 전", "변경 후"],
    [
        ["backend/embedder.py", "qwen3-embedding:4b", "qwen3-embedding:0.6b"],
        ["backend/rag.py",      "mxbai-embed-large",  "qwen3-embedding:0.6b  +  [:1024] 슬라이싱 추가"],
    ],
    col_widths=[4.5, 4.0, 6.0]
)

add_action_label(doc, "🟡 참고 사항  —  backend/save_articles.py", urgent=False)
add_body(doc, "UNVERIFIED 기사 저장 시 needs_review=True가 자동으로 설정됩니다. 별도 코드 수정 없이 동작합니다.")

add_action_label(doc, "🟡 참고 사항  —  backend/main.py 신규 엔드포인트", urgent=False)
add_table(doc,
    ["메서드", "경로", "역할"],
    [
        ["GET",   "/admin/review",            "needs_review=TRUE 기사 목록 반환 (최대 50건)"],
        ["PATCH", "/admin/review/{url_hash}", "관리자 판정 입력 → fact_label 업데이트 + needs_review=FALSE"],
    ],
    col_widths=[2.0, 5.5, 7.0]
)
add_body(doc, "PATCH 요청 본문 예시:")
add_code_block(doc, [
    '{ "verdict": "FACT",  "reviewer_note": "공식 발표 확인됨" }',
    '{ "verdict": "RUMOR", "reviewer_note": "출처 불명확, 수치 과장" }',
])


# ══════════════════════════════════════════════════════════════
# 정수민
# ══════════════════════════════════════════════════════════════
add_person_heading(doc, "정수민", "프론트엔드 UI/UX", "2E7D32")

add_action_label(doc, "🔴 PM과 협의 필요  —  UNVERIFIED 기사 UI 처리 방식", urgent=True)
add_body(doc, "팩트체크 결과가 UNVERIFIED인 기사에 needs_review=true 플래그가 붙습니다. 아래 두 가지 방식 중 하나를 김민규님과 결정해주세요.")
add_table(doc,
    ["선택지", "설명", "장단점"],
    [
        ["A) 피드에서 숨김",
         "needs_review=TRUE 기사는\n관리자 판정 완료 전까지 노출 안 함",
         "✅ 사용자에게 미검증 정보 노출 없음\n❌ 기사 수가 줄어들 수 있음"],
        ["B) '검토 중' 뱃지 표시",
         "fact_label 뱃지를 'UNVERIFIED' 대신\n'검토 중'으로 표시하고 피드에 노출",
         "✅ 기사 수 유지\n❌ 사용자가 미검증 정보 볼 수 있음"],
    ],
    col_widths=[3.0, 5.5, 6.0]
)
add_body(doc, "API 응답에 needs_review 필드가 포함되어 있으므로 결정되면 해당 필드로 분기 처리하면 됩니다.")

add_action_label(doc, "🟡 참고 사항  —  어드민 검토 화면 (선택)", urgent=False)
add_body(doc, "어드민 화면이 필요할 경우 아래 API로 연동 가능합니다.")
add_table(doc,
    ["API", "용도"],
    [
        ["GET  /admin/review",            "검토 대기 기사 목록 불러오기"],
        ["PATCH /admin/review/{url_hash}", "판정 버튼(FACT / RUMOR / 무시) 클릭 시 호출"],
    ],
    col_widths=[5.5, 9.0]
)


# ══════════════════════════════════════════════════════════════
# 이상준
# ══════════════════════════════════════════════════════════════
add_person_heading(doc, "이상준", "RSS 크롤러 + 데이터 전처리", "6A1B9A")

add_action_label(doc, "🔴 확인 필요  —  source 필드값 일치 여부", urgent=True)
add_body(doc, "팩트체크 파이프라인이 source 이름을 기준으로 채널 신뢰도를 조회합니다. RSS 크롤러에서 넘기는 source 값이 아래 목록과 정확히 일치해야 합니다.")
add_table(doc,
    ["channel_config.py 등록 이름", "source_type", "기본 tier", "credibility_score"],
    [
        ["MIT Technology Review",  "media",     "MEDIA_OFFICIAL",        "0.95"],
        ["IEEE Spectrum",          "media",     "MEDIA_OFFICIAL",        "0.93"],
        ["The Guardian Tech",      "media",     "MEDIA_OFFICIAL",        "0.88"],
        ["TechCrunch",             "media",     "MEDIA_OFFICIAL",        "0.82"],
        ["The Verge",              "media",     "MEDIA_OFFICIAL",        "0.80"],
        ["The Decoder",            "media",     "MEDIA_OFFICIAL",        "0.82"],
        ["VentureBeat AI",         "media",     "MEDIA_CREDIBLE_LEAK",   "0.72"],
    ],
    col_widths=[4.5, 2.5, 4.0, 3.5]
)
add_callout(doc,
    "source 이름이 다르면 (예: 'Venturebeat' vs 'VentureBeat AI') 미등록 출처로 처리되어 credibility_score 0.50이 적용됩니다.",
    bg="FDECEA", color=(180, 0, 0)
)

add_action_label(doc, "🟡 참고 사항  —  파일 변경", urgent=False)
add_table(doc,
    ["파일", "변경 내용"],
    [
        ["collect/classifier.py",        "신규 추가 (GitHub 머지) — 크롤러 파이프라인 연동 여부 확인 필요"],
        ["collect/crawler/rss_crawler.py", "수정됨 (GitHub 머지) — 변경 내용 확인 권장"],
        ["pipeline/translator.py",       "삭제됨 — pipeline/translate_summarize.py로 통합"],
        ["pipeline/summarizer.py",       "삭제됨 — pipeline/translate_summarize.py로 통합"],
    ],
    col_widths=[5.0, 9.5]
)


# ══════════════════════════════════════════════════════════════
# 김민규 (PM)
# ══════════════════════════════════════════════════════════════
add_person_heading(doc, "김민규", "PM + 프론트엔드 총괄", "B71C1C")

add_action_label(doc, "🔴 정책 결정 필요  —  UNVERIFIED 기사 피드 노출 방식", urgent=True)
add_body(doc, "정수민 파트 참고. A(숨김) / B(검토 중 뱃지) 중 결정 후 정수민에게 전달해주세요.")

add_action_label(doc, "🔴 인지 필요  —  VentureBeat API 비용 증가", urgent=True)
add_body(doc, "VentureBeat AI 출처 기사는 이제 루머 신호가 없어도 Gemini API 검증을 거칩니다 (기존: Step 1 FACT_AUTO 즉시 처리).")
add_body(doc, "VentureBeat 일 평균 기사 수 × Gemini API 단가만큼 비용 증가 예상", indent=0.5, bullet=True)
add_body(doc, "소식통 인용 출처의 팩트체크 정확도를 위한 의도적 트레이드오프", indent=0.5, bullet=True)

add_action_label(doc, "🟡 전체 변경 요약", urgent=False)
add_table(doc,
    ["항목", "변경 전", "변경 후", "효과"],
    [
        ["임베딩 모델",        "파일마다 다름\n(mxbai, 4B 혼재)",    "Qwen3-Embedding-0.6B\n전체 통일",        "일관성 확보"],
        ["Gemini Pass A",     "LLM으로 문체 분석",                   "제거 → signal_detector\n결과 재사용",    "API 1회 절감"],
        ["CoVe Draft",        "별도 LLM 호출",                       "Step 3A 결과 재사용",                   "API 1회 절감"],
        ["CREDIBLE_LEAK 처리", "루머 신호 없으면\nFACT_AUTO 즉시",    "Gemini 검증 강제",                      "정확도 향상"],
        ["UNVERIFIED 처리",   "DB에 저장 후 방치",                   "needs_review 플래그\n+ 어드민 판정 구조", "Human-in-the-Loop"],
    ],
    col_widths=[3.5, 3.5, 3.5, 4.0]
)

add_divider(doc)
p_end = doc.add_paragraph()
p_end.paragraph_format.space_before = Pt(10)
r_end = p_end.add_run("작성: 이동우  ·  2026-04-21  |  상세 내용: 팩트체크_파이프라인_개선_팀원공유.docx 참고")
set_font(r_end, size=9, color=(130, 130, 130))


# ── 저장 ────────────────────────────────────────────────────
out_path = r"C:\Users\이동우\samseon\팀원_공유사항_담당별정리.docx"
doc.save(out_path)
print(f"저장 완료: {out_path}")
