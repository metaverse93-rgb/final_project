"""
팩트체크 파이프라인 개선 내역 팀원 공유 문서
실행: python make_factcheck_update_doc.py
출력: 팩트체크_파이프라인_개선_팀원공유.docx
"""

import sys
sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(3.0)


# ── 스타일 헬퍼 ─────────────────────────────────────────────
def set_font(run, size=11, bold=False, color=None, name="맑은 고딕"):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), name)
    rPr.insert(0, rFonts)

def add_heading(doc, text, level=1, color=(0, 70, 127)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    sizes = {1: 16, 2: 13, 3: 11}
    set_font(run, size=sizes.get(level, 11), bold=True, color=color)
    return p

def add_body(doc, text, indent=0, bullet=False, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    prefix = "• " if bullet else ""
    run = p.add_run(prefix + text)
    set_font(run, size=10)
    return p

def add_bold_body(doc, label, text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    r1 = p.add_run(label)
    set_font(r1, size=10, bold=True, color=(0, 70, 127))
    r2 = p.add_run(text)
    set_font(r2, size=10)
    return p

def shade_cell(cell, hex_color="DEEAF1"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, size=10, color=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)

def add_table(doc, headers, rows, col_widths=None, header_color="1F4E79"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        shade_cell(cell, header_color)
        set_cell_text(cell, h, bold=True, size=10, color=(255,255,255), align="center")
    for i, row in enumerate(rows):
        bg = "F2F2F2" if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(row):
            cell = table.rows[i+1].cells[j]
            shade_cell(cell, bg)
            set_cell_text(cell, val, size=9)
    if col_widths:
        for row in table.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run("─" * 55)
    set_font(run, size=9, color=(180, 180, 180))

def add_code_block(doc, lines, bg="F0F0F0"):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.5)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.space_before = Pt(1)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  bg)
        pPr.append(shd)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)

def add_callout(doc, text, bg="FFF2CC", color=(127, 96, 0)):
    """강조 박스 (노란색 배경 등)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  bg)
    pPr.append(shd)
    run = p.add_run(text)
    set_font(run, size=10, color=color)
    return p


# ══════════════════════════════════════════════════════════════
# 표지
# ══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run("삼선뉴스 (三鮮)")
set_font(run, size=22, bold=True, color=(0, 70, 127))

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("팩트체크 파이프라인 개선 내역\n팀원 공유 문서")
set_font(run2, size=14, color=(80, 80, 80))

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(8)
run3 = p3.add_run("담당: 이동우 · 2026-04-21  |  6주차  (v2 — CREDIBLE_LEAK 라우팅 + Human-in-the-Loop 추가)")
set_font(run3, size=10, color=(130, 130, 130))

doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# 1. 변경 배경
# ══════════════════════════════════════════════════════════════
add_heading(doc, "1. 왜 바꿨나요?", level=1)
add_body(doc, "기획서에서 설계한 3단계 멀티에이전트 팩트체크 파이프라인(Step 3A → 3B → 3C)을 구현하면서, 기존 방식과 멀티에이전트 방식 사이에 중복되는 부분이 발견되었습니다. 이를 정리하고 더 효율적인 구조로 개선했습니다.")

add_body(doc, "중복 제거: Gemini Pass A(문체 분석)가 Step 1 signal_detector(정규식 패턴 탐지)와 동일한 작업을 수행 → Pass A 제거 후 signal_detector 결과를 Gemini에 직접 전달", indent=0.8, bullet=True)
add_body(doc, "API 호출 절감: CoVe(Step 3B)에서 별도 Draft 생성 호출을 없애고 Step 3A 결과를 재사용 → 기사 1건당 최대 3회 API 호출 감소", indent=0.8, bullet=True)
add_body(doc, "계층 구조 유지: Step 0 → 1 → 2 → 3A → 3B → 3C 흐름은 그대로 유지, 라우팅 로직 변경 없음", indent=0.8, bullet=True)

add_callout(doc,
    "핵심 원칙: 계층 구조(비용 절감 목적)는 그대로 유지하면서, 꼭 필요한 기사에만 멀티에이전트(DebateCV) 판정이 투입됩니다.",
    bg="E7F3FF", color=(0, 70, 127)
)


# ══════════════════════════════════════════════════════════════
# 2. 파이프라인 전체 흐름 비교
# ══════════════════════════════════════════════════════════════
add_heading(doc, "2. 파이프라인 전체 흐름 비교", level=1)

add_heading(doc, "2-1. 변경 전 (기획서 초기 구현)", level=2)
add_code_block(doc, [
    "기사 입력",
    "  │",
    "  ▼ Step 0: 채널 등급 분류 (channel_config.py)        ← 규칙 기반, 무료",
    "  │",
    "  ▼ Step 1: 루머 신호 패턴 탐지 (signal_detector.py)  ← 규칙 기반, 무료",
    "  │",
    "  ▼ Step 2: Google Fact Check API (~15.8% 매칭)        ← 무료",
    "  │",
    "  ▼ Step 3A: Gemini 2-Pass Advisor                     ← API 2회 호출",
    "  │   ├─ Pass A: 문체 분석 (감정어/헤징/출처 스타일)  ← [중복] signal_detector와 겹침",
    "  │   └─ Pass B: Google Search Grounding + 추론",
    "  │",
    "  ▼ Step 3B: CoVe 자기검증 (confidence < 0.80 시)      ← API 5회 호출",
    "  │   ├─ Step 1: Draft verdict 생성                    ← [중복] Step 3A 결과와 겹침",
    "  │   ├─ Step 2: 검증 질문 5개 생성",
    "  │   ├─ Step 3: 각 질문 독립 답변 (×3)",
    "  │   └─ Step 4: 종합 판단",
    "  │",
    "  ▼ Step 3C: DebateCV 3-에이전트 (importance ≥ 0.7 시) ← API 5회 호출",
    "      ├─ Round 1: Prosecutor + Defender 논거 생성",
    "      ├─ Round 2: 상호 반박",
    "      └─ Judge: 최종 판결",
])

add_heading(doc, "2-2. 변경 후 (개선 버전)", level=2)
add_code_block(doc, [
    "기사 입력",
    "  │",
    "  ▼ Step 0: 채널 등급 분류 (channel_config.py)        ← 규칙 기반, 무료  [변경 없음]",
    "  │",
    "  ▼ Step 1: 루머 신호 패턴 탐지 (signal_detector.py)  ← 규칙 기반, 무료  [변경 없음]",
    "  │         └─ 결과(신호 강도, 매칭 패턴) → Step 3A로 전달",
    "  │",
    "  ▼ Step 2: Google Fact Check API (~15.8% 매칭)        ← 무료             [변경 없음]",
    "  │",
    "  ▼ Step 3A: Gemini Advisor                            ← API 1회 호출  ★ 1회 절감",
    "  │   └─ signal_detector 결과 + Google Search Grounding + 추론",
    "  │      (Pass A 문체 분석 제거 — signal_detector 결과로 대체)",
    "  │",
    "  ▼ Step 3B: CoVe 자기검증 (confidence < 0.80 시)      ← API 4회 호출  ★ 1회 절감",
    "  │   ├─ Draft: Step 3A 결과 재사용 (별도 호출 없음)  ← Draft 호출 제거",
    "  │   ├─ Step 2: 검증 질문 5개 생성",
    "  │   ├─ Step 3: 각 질문 독립 답변 (×3)",
    "  │   └─ Step 4: 종합 판단",
    "  │",
    "  ▼ Step 3C: DebateCV 3-에이전트 (importance ≥ 0.7 시) ← API 5회 호출  [변경 없음]",
    "      ├─ Round 1: Prosecutor + Defender 논거 생성",
    "      ├─ Round 2: 상호 반박",
    "      └─ Judge: 최종 판결",
])


# ══════════════════════════════════════════════════════════════
# 3. 핵심 변경 상세
# ══════════════════════════════════════════════════════════════
add_heading(doc, "3. 핵심 변경 상세", level=1)

add_heading(doc, "3-1. Pass A 제거 → signal_detector 결과 재사용", level=2)
add_body(doc, "Pass A가 LLM으로 분석하던 항목들이 signal_detector.py의 정규식 패턴과 동일했습니다.")

add_table(doc,
    ["Pass A가 하던 일", "signal_detector가 이미 하던 일", "결론"],
    [
        ["감정어/과장 표현 탐지", "OPINION_PATTERNS (opinion, editorial, I think 등)", "중복"],
        ["헤징 표현 밀도 측정", "RUMOR_WEAK_PATTERNS (possibly, expected to, 추측 등)", "중복"],
        ["출처 인용 방식 분석", "CREDIBLE_LEAK_PATTERNS (sources say, 단독, 소식통 등)", "중복"],
        ["헤드라인 과장 여부", "RUMOR_STRONG_PATTERNS (allegedly, reportedly 등)", "중복"],
    ],
    col_widths=[5.0, 7.0, 2.5]
)

add_body(doc, "변경 후: signal_detector 결과(신호 강도 STRONG/WEAK/NONE, 매칭 패턴 목록)를 Gemini 프롬프트에 직접 주입합니다.")

add_code_block(doc, [
    "# 변경 전",
    "gemini_result = gemini_run(title, content)",
    "",
    "# 변경 후",
    "gemini_result = gemini_run(",
    "    title, content,",
    "    signal_strength=signal.rumor_strength,   # 'STRONG' / 'WEAK' / 'NONE'",
    "    signal_patterns=signal.matched_patterns, # ['allegedly', 'sources say', ...]",
    ")",
])

add_heading(doc, "3-2. CoVe Draft 호출 제거 → Step 3A 결과 재사용", level=2)
add_body(doc, "CoVe(Chain-of-Verification)는 원래 4단계로 동작했습니다.")

add_table(doc,
    ["단계", "변경 전", "변경 후", "API 호출"],
    [
        ["Step 1 (Draft)", "Gemini 호출 → 새 Draft verdict 생성", "Step 3A 결과(prior_verdict) 그대로 사용", "1회 → 0회"],
        ["Step 2 (질문 생성)", "검증 질문 5개 독립 생성", "동일 (변경 없음)", "1회"],
        ["Step 3 (독립 답변)", "각 질문 독립 답변 ×3", "동일 (변경 없음)", "3회"],
        ["Step 4 (종합)", "Draft vs 검증 결과 비교 → 최종 판단", "동일 (변경 없음)", "1회"],
    ],
    col_widths=[3.0, 5.5, 5.5, 2.5]
)

add_callout(doc,
    "CoVe 핵심 원칙은 유지됩니다: '검증 질문 생성 시 Draft를 절대 보여주지 않는다.' — Step 2~4에서 이 원칙은 그대로 지켜집니다. Step 3A 결과를 Draft로 쓰더라도 질문 생성 프롬프트에 Draft가 포함되지 않으므로 논문 원칙과 동일합니다.",
    bg="E8F5E9", color=(27, 94, 32)
)


# ══════════════════════════════════════════════════════════════
# 4. API 호출 횟수 변화
# ══════════════════════════════════════════════════════════════
add_heading(doc, "4. API 호출 횟수 변화", level=1)
add_body(doc, "기사 유형별로 실제 API 호출 횟수가 얼마나 달라지는지 정리합니다.")

add_table(doc,
    ["기사 유형", "도달 단계", "변경 전 API 호출", "변경 후 API 호출", "절감"],
    [
        ["공식 미디어 + 루머 신호 없음\n(MIT TR, IEEE 등)", "Step 1 종료\n(FACT_AUTO)", "0회", "0회", "-"],
        ["Google FC API 매칭\n(~15.8%)", "Step 2 종료", "0회", "0회", "-"],
        ["일반 미매칭 기사\n(confidence ≥ 0.80)", "Step 3A 종료", "2회", "1회", "1회 ↓"],
        ["불확실 기사\n(confidence < 0.80,\nimportance < 0.70)", "Step 3B 종료", "7회", "5회", "2회 ↓"],
        ["고중요도 불확실 기사\n(모델명+수치+SOTA 포함)", "Step 3C 종료", "12회", "10회", "2회 ↓"],
    ],
    col_widths=[4.5, 3.0, 3.0, 3.0, 2.0]
)

add_body(doc, "멀티에이전트(DebateCV) 발동 조건은 변경 없습니다: importance_score ≥ 0.70 AND confidence < 0.80", space_after=2)
add_body(doc, "importance_score 계산 기준: AI 모델명 포함(×0.4) + 벤치마크 수치(×0.3) + SOTA/최초 주장(×0.2) + Breaking 뉴스(×0.1)")


# ══════════════════════════════════════════════════════════════
# 5. 멀티에이전트 판정 방식 설명
# ══════════════════════════════════════════════════════════════
add_heading(doc, "5. 멀티에이전트 판정 방식이란? (Step 3C)", level=1)
add_body(doc, "DebateCV는 팩트체크가 특히 어려운 '고중요도 기사'에만 투입되는 3인 토론 시스템입니다.")

add_heading(doc, "5-1. 왜 단일 LLM 판단이 부족한가?", level=2)
add_body(doc, "\"GPT-5가 MMLU 97% 달성, 세계 최초\" 같은 기사는 단일 LLM이 판단하면 편향이 생깁니다.")
add_body(doc, "LLM은 자신의 학습 데이터와 비슷한 주장을 사실로 믿는 경향이 있음", indent=0.8, bullet=True)
add_body(doc, "반대 의견을 스스로 생성하지 않으면 한쪽으로 치우친 판단을 내림", indent=0.8, bullet=True)
add_body(doc, "해결책: 서로 다른 역할을 맡은 에이전트들이 강제로 반대 입장을 취하게 함", indent=0.8, bullet=True)

add_heading(doc, "5-2. 3인 토론 구조", level=2)

add_table(doc,
    ["역할", "맡은 일", "Temperature", "근거 논문"],
    [
        ["Prosecutor\n(검사)", "\"이 기사가 왜 거짓/과장인가\"\n기술적 오류, 수치 과장, 출처 불명 등 약점 공격", "0.7\n(다양한 논거)", "Chern et al.\narXiv:2507.19090"],
        ["Defender\n(변호인)", "\"이 기사가 왜 사실인가\"\n공식 발표 일치, 수치 타당성, 맥락 공정성 방어", "0.7\n(다양한 논거)", "Chern et al.\narXiv:2507.19090"],
        ["Judge\n(판사)", "양측 논거 + 반박 검토\n→ 최종 verdict + confidence 결정", "0.3\n(신중한 판단)", "Du et al.\nICML 2024"],
    ],
    col_widths=[2.5, 6.5, 2.0, 3.5]
)

add_heading(doc, "5-3. 토론 흐름", level=2)
add_code_block(doc, [
    "Round 1: Prosecutor 논거 생성  +  Defender 논거 생성  (동시, 독립적으로)",
    "          ↓                           ↓",
    "Round 2: Prosecutor → Defender 논거 반박",
    "         Defender   → Prosecutor 논거 반박",
    "          ↓",
    "Final:   Judge → 양측 논거 + 반박 전체 검토 → 최종 verdict + confidence",
    "",
    "만장일치(UNANIMOUS) 시: confidence +0.05 보정",
    "2:1(MAJORITY) 시: confidence 0.70~0.80",
    "팽팽(SPLIT) 시: UNVERIFIED 처리",
])

add_heading(doc, "5-4. 기존 방식 vs 멀티에이전트 판정 비교", level=2)

add_table(doc,
    ["구분", "기존 방식 (Step 3A/3B)", "멀티에이전트 판정 (Step 3C)"],
    [
        ["에이전트 수", "1개 (같은 모델 반복 호출)", "3개 (Prosecutor / Defender / Judge)"],
        ["판정 방식", "단일 관점에서 분석 후 판단", "서로 반대 입장에서 토론 후 판사가 결정"],
        ["편향 제어", "자기 편향 있음", "강제 반박으로 편향 상쇄"],
        ["투입 조건", "모든 미매칭 기사", "importance ≥ 0.70 고중요도 기사만"],
        ["API 호출", "1~4회", "5회"],
        ["정확도 향상", "기준선", "단일 LLM 대비 F1 +9.2%\n(Chern et al. 2024)"],
    ],
    col_widths=[3.0, 5.5, 5.5]
)


# ══════════════════════════════════════════════════════════════
# 6. 추가 개선: CREDIBLE_LEAK 라우팅 수정
# ══════════════════════════════════════════════════════════════
add_heading(doc, "6. 추가 개선 ①: CREDIBLE_LEAK 출처 자동 FACT 방지", level=1)
add_body(doc, "시뮬레이션 중 발견된 문제: VentureBeat처럼 기본 등급이 MEDIA_CREDIBLE_LEAK인 출처가 루머 신호가 없을 때 Step 1에서 바로 FACT_AUTO로 처리되고 있었습니다.")

add_heading(doc, "6-1. 문제 원인", level=2)
add_code_block(doc, [
    "# 기존 코드 (pipeline.py)",
    "if signal.fact_label_hint == 'FACT_AUTO':   # tier 무관하게 즉시 FACT 반환",
    "    return FactCheckResult(fact_label='FACT', ...)",
    "",
    "# VentureBeat의 경우:",
    "# - default_tier = MEDIA_CREDIBLE_LEAK  (소식통 인용이 잦은 출처)",
    "# - 루머 신호 패턴 없음  →  signal_detector가 FACT_AUTO 힌트 반환",
    "# - 결과: 검증 없이 FACT 처리  ← 문제!",
])

add_heading(doc, "6-2. 수정 내용", level=2)
add_code_block(doc, [
    "# 수정 후 코드 (pipeline.py)",
    "# MEDIA_CREDIBLE_LEAK는 루머 신호가 없어도 LLM 검증 필요",
    "if signal.fact_label_hint == 'FACT_AUTO' and tier != 'MEDIA_CREDIBLE_LEAK':",
    "    return FactCheckResult(fact_label='FACT', ...)",
    "",
    "# CREDIBLE_LEAK tier는 조건을 통과하지 못하고 Step 2/3으로 계속 진행됨",
])

add_table(doc,
    ["출처 예시", "tier", "수정 전", "수정 후"],
    [
        ["MIT Technology Review", "MEDIA_OFFICIAL", "Step 1 → FACT_AUTO ✅", "Step 1 → FACT_AUTO ✅ (변경 없음)"],
        ["VentureBeat AI", "MEDIA_CREDIBLE_LEAK", "Step 1 → FACT_AUTO ❌", "Step 3A Gemini 검증 ✅"],
        ["The Decoder", "MEDIA_OFFICIAL\n(Leak 신호 있을 때 CREDIBLE_LEAK으로 조정)", "신호 있으면 이미 Step 3 진행", "동일 (변경 없음)"],
    ],
    col_widths=[3.8, 3.5, 3.5, 3.7]
)

add_callout(doc,
    "영향: VentureBeat AI는 매일 수십 건의 기사를 발행하므로 이 수정으로 해당 출처의 모든 기사가 Gemini 검증을 받게 됩니다. API 호출 증가가 예상되나 소식통 인용 출처의 정확도를 위해 필요한 트레이드오프입니다.",
    bg="FFF2CC", color=(127, 96, 0)
)


# ══════════════════════════════════════════════════════════════
# 7. 추가 개선: Human-in-the-Loop (UNVERIFIED 수동 검토)
# ══════════════════════════════════════════════════════════════
add_heading(doc, "7. 추가 개선 ②: UNVERIFIED → Human-in-the-Loop", level=1)
add_body(doc, "DebateCV 3인 토론에서 SPLIT(팽팽) 판정이 나거나 LLM이 판단을 포기하면 UNVERIFIED 라벨이 붙습니다. 이 기사를 자동 처리하지 않고 사람이 최종 판정하도록 구조를 추가했습니다.")

add_heading(doc, "7-1. UNVERIFIED가 발생하는 경우", level=2)
add_table(doc,
    ["상황", "confidence", "처리"],
    [
        ["DebateCV 3인 토론에서 SPLIT 판정", "0.50~0.65", "UNVERIFIED → 수동 검토 대기"],
        ["Gemini/CoVe API 예외 발생 + importance 낮음", "0.50", "UNVERIFIED → 수동 검토 대기"],
        ["Google FC API 미매칭 + skip_llm=True (테스트 모드)", "0.50", "UNVERIFIED → 수동 검토 대기"],
    ],
    col_widths=[7.0, 2.5, 4.5]
)

add_heading(doc, "7-2. 처리 흐름", level=2)
add_code_block(doc, [
    "파이프라인 결과 → fact_label = 'UNVERIFIED'",
    "                           ↓",
    "  DB 저장 시 needs_review = TRUE  (articles 테이블 신규 컬럼)",
    "                           ↓",
    "  GET /admin/review         ← 관리자가 대기 목록 조회",
    "                           ↓",
    "  [FACT 확인] / [RUMOR 확인] / [무시] 버튼 클릭",
    "                           ↓",
    "  PATCH /admin/review/{url_hash}  { verdict: 'FACT' | 'RUMOR' | 'UNVERIFIED' }",
    "                           ↓",
    "  articles.fact_label 업데이트 + needs_review = FALSE",
    "  fact_checks 테이블에 checker_type='human' 기록 저장",
])

add_heading(doc, "7-3. 추가된 API 엔드포인트", level=2)
add_table(doc,
    ["메서드", "경로", "역할", "담당"],
    [
        ["GET",   "/admin/review",              "needs_review=TRUE 기사 목록 반환", "이동우 (backend)"],
        ["PATCH", "/admin/review/{url_hash}",   "verdict 직접 입력 → fact_label 업데이트", "이동우 (backend)"],
    ],
    col_widths=[2.0, 5.0, 5.5, 2.5]
)

add_heading(doc, "7-4. DB 변경 사항", level=2)
add_code_block(doc, [
    "-- articles 테이블 신규 컬럼 (supabase_schema.sql 반영)",
    "needs_review  BOOLEAN DEFAULT FALSE,",
    "-- TRUE  → 관리자 수동 검토 대기 (UNVERIFIED 기사)",
    "-- FALSE → 검토 완료 또는 검토 불필요",
    "",
    "-- fact_checks 테이블 — 기존 checker_type 컬럼 활용",
    "checker_type = 'human'  -- 관리자가 직접 판정한 기록",
])

add_callout(doc,
    "프론트엔드 연동: UNVERIFIED 기사는 사용자 피드에 노출되지만 fact_label 뱃지가 '검토 중'으로 표시되도록 정수민님과 협의가 필요합니다. needs_review=TRUE인 기사를 피드에서 숨길지 여부도 PM(김민규님)이 결정해주세요.",
    bg="FCE4EC", color=(136, 14, 79)
)


# ══════════════════════════════════════════════════════════════
# (기존 6번 → 8번으로 변경)
# ══════════════════════════════════════════════════════════════
add_heading(doc, "8. 수정된 파일 목록", level=1)

add_table(doc,
    ["파일", "변경 내용", "팀원 영향"],
    [
        ["fact_checker/gemini_advisor.py",
         "Pass A 제거\nrun() 시그니처에 signal_strength, signal_patterns 추가",
         "없음 (pipeline.py가 자동 처리)"],
        ["fact_checker/cove_verifier.py",
         "Draft LLM 호출 제거\nprior_verdict를 Draft로 재사용",
         "없음"],
        ["fact_checker/pipeline.py",
         "gemini_run() 호출 시 signal 결과 전달\nCREDIBLE_LEAK FACT_AUTO 방지 조건 추가",
         "없음"],
        ["fact_checker/signal_detector.py",
         "주석 업데이트",
         "없음"],
        ["fact_checker/channel_config.py",
         "주석 업데이트",
         "없음"],
        ["backend/save_articles.py",
         "UNVERIFIED 저장 시 needs_review=True 자동 설정",
         "없음 (자동 처리)"],
        ["backend/main.py",
         "GET /admin/review 엔드포인트 추가\nPATCH /admin/review/{url_hash} 엔드포인트 추가",
         "정수민(프론트): 어드민 UI 연동 시 참고"],
        ["supabase_schema.sql",
         "articles 테이블에 needs_review BOOLEAN 컬럼 추가",
         "강주찬: Supabase SQL Editor에서 ALTER TABLE 실행 필요"],
    ],
    col_widths=[4.5, 6.0, 4.0]
)

add_callout(doc,
    "팀원 연동 포인트는 변경 없습니다. save_articles.py → run_fact_check() → DB 저장 흐름은 동일합니다.",
    bg="FFF2CC", color=(127, 96, 0)
)


# ══════════════════════════════════════════════════════════════
# 7. 현재 KPI 목표
# ══════════════════════════════════════════════════════════════
add_heading(doc, "9. 팩트체크 KPI 목표", level=1)

add_table(doc,
    ["지표", "목표", "현재 상태", "측정 시점"],
    [
        ["신뢰도 분류 정확도", "≥ 80%", "미측정", "6주차 수동 100건 검증"],
        ["RUMOR recall", "≥ 0.75", "미측정", "6주차"],
        ["처리 속도 (일반 기사)", "≤ 5초/건", "미측정", "통합 테스트 시"],
        ["처리 속도 (DebateCV)", "≤ 30초/건", "미측정", "통합 테스트 시"],
        ["Google FC 매칭률", "~15.8%", "AI 테크 도메인 특성", "고정값"],
    ],
    col_widths=[4.5, 2.5, 3.0, 4.5]
)

add_divider(doc)

p_end = doc.add_paragraph()
p_end.paragraph_format.space_before = Pt(12)
r_end = p_end.add_run("문의: 이동우 (Discord / KakaoTalk)  |  코드 위치: fact_checker/")
set_font(r_end, size=9, color=(130, 130, 130))


# ── 저장 ────────────────────────────────────────────────────
out_path = r"C:\Users\이동우\samseon\팩트체크_파이프라인_개선_팀원공유.docx"
doc.save(out_path)
print(f"저장 완료: {out_path}")
