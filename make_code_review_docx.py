"""
eval/CODE_REVIEW_0417.md → DOCX 변환
출력: Desktop/삼선뉴스_코드리뷰_0417.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUTPUT = os.path.join(os.path.expanduser("~"), "Desktop", "삼선뉴스_코드리뷰_0417.docx")

doc = Document()

# ── 기본 스타일 ──────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "맑은 고딕"
style.font.size = Pt(10.5)

def heading1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

def heading2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def heading3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = RGBColor(0xC0, 0x50, 0x00)

def body(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "맑은 고딕"
    run.font.size = Pt(10.5)
    if bold:
        run.bold = True
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x17, 0x17, 0x17)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    shading_elm = p._element
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = shading_elm.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    return p

def bullet(text, indent=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + indent * 0.25)
    run = p.add_run(text)
    run.font.name = "맑은 고딕"
    run.font.size = Pt(10.5)

def add_table(headers, rows_data):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.name = "맑은 고딕"
            run.font.size = Pt(10)
    for row_data in rows_data:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
            for run in row[i].paragraphs[0].runs:
                run.font.name = "맑은 고딕"
                run.font.size = Pt(10)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════
# 표지
# ══════════════════════════════════════════════════════
doc.add_heading("파이프라인 코드 리뷰 보고서", level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("삼선뉴스 프로젝트 — 2026-04-17\n이동우 (번역·요약 파이프라인 담당)")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
doc.add_paragraph()

# ══════════════════════════════════════════════════════
# 1. 리뷰 범위
# ══════════════════════════════════════════════════════
heading1("1. 리뷰 범위")
files = [
    "eval/run_eval_finetuned.py",
    "eval/run_eval_base.py",
    "eval/metrics/term_preservation.py",
    "eval/metrics/geval.py",
    "eval/metrics/bleu_comet.py",
    "pipeline/translator.py",
    "pipeline/summarizer.py",
    "pipeline/translate_summarize.py",
    "backend/save_articles.py",
]
for f in files:
    bullet(f)
doc.add_paragraph()

# ══════════════════════════════════════════════════════
# 2. CRITICAL 이슈
# ══════════════════════════════════════════════════════
heading1("2. CRITICAL — 점수에 직접 영향")

# Issue 1
heading2("Issue 1: geval.py Fluency 기준이 guide v2와 정반대")
body("파일: eval/metrics/geval.py:82-94", bold=True)
doc.add_paragraph()
body("문제")
body("G-Eval Fluency 채점 기준이 구 규칙(음차 허용)으로 작성되어 있음.")
code_block("② Standard transliterations: fine-tuning→파인튜닝, embedding→임베딩\n③ Proper nouns FIRST mention: EnglishName(한국어 음차), e.g., OpenAI(오픈에이아이)")
doc.add_paragraph()
body("충돌")
body("파인튜닝 모델의 SUMMARIZE_SYSTEM은 guide v2 기준으로 영문 유지:")
code_block("고유명사(회사명·제품명·인물명)는 영문 그대로 유지 (Nvidia, OpenAI ...)")
doc.add_paragraph()
body("영향")
body('파인튜닝 모델이 "Fine-tuning", "Nvidia"를 올바르게 출력해도 G-Eval 평가자가 -1점 패널티 부과 → G-Eval 점수 억제 (현재 3.64, 목표 4.0 미달의 원인 중 하나).')
doc.add_paragraph()
body("수정 내용")
body("Fluency 기준 ②③을 guide v2 기준으로 교체:")
code_block("② AI/기술 용어는 영문 그대로: Fine-tuning, Embedding, Prompt (파인튜닝, 임베딩 사용 시 -1점)\n③ 고유명사는 영문만 사용: Nvidia, OpenAI (한국어 음차 사용 시 -1점)")
doc.add_paragraph()

# Issue 2
heading2("Issue 2: translate_summarize.py (프로덕션)가 guide v2와 충돌")
body("파일: pipeline/translate_summarize.py:53-68", bold=True)
doc.add_paragraph()
body("문제")
body("Rule 3·4가 여전히 음차 규칙 사용:")
code_block("# Rule 3\nFine-tuning→파인튜닝 / Embedding→임베딩 / Prompt→프롬프트\n\n# Rule 4\nEnglishName(한국어 음차) on FIRST mention only\ne.g., Anthropic(앤트로픽) / OpenAI(오픈에이아이) / Nvidia(엔비디아)")
doc.add_paragraph()
body("영향")
bullet("run_eval_base.py가 이 함수 import → 베이스라인 평가 출력에 음차 포함 → TPR 낮게 측정")
bullet("프로덕션 서비스 출력이 guide v2 불일치 상태")
bullet("파인튜닝 데이터(guide v2 기준)와 프로덕션 추론 결과의 train/inference mismatch")
doc.add_paragraph()
body("수정 내용")
body("Rule 3: 음차 지시 → 영문 유지 지시로 변경. Rule 4: 한국어 음차 병기 제거. Rule 5: 인물명 음차 제거.")
doc.add_paragraph()

# ══════════════════════════════════════════════════════
# 3. MEDIUM 이슈
# ══════════════════════════════════════════════════════
heading1("3. MEDIUM — TPR/평가 정확도 영향")

heading2("Issue 3: print_summary() mean()이 TPR=0을 유효 데이터에서 제외")
body("파일: eval/run_eval_base.py:112, eval/run_eval_finetuned.py 동일", bold=True)
doc.add_paragraph()
body("문제")
code_block('vals = [float(r[key]) for r in rows if r.get(key) not in ("", "0", "0.0", None)]')
body('TPR=0(용어 완전 미보존)은 실제 발생 가능한 유효한 값인데 평균 계산에서 제외됨.')
doc.add_paragraph()
body("영향")
body("보고 TPR이 실제보다 높게 나옴. 현재 85.9%가 실제보다 과대 추정일 수 있음.")
doc.add_paragraph()
body("수정 내용")
code_block('# TPR만 선택적으로 0 포함\nif key == "tpr":\n    vals = [float(r[key]) for r in rows if r.get(key) not in ("", None)]\nelse:\n    vals = [float(r[key]) for r in rows if r.get(key) not in ("", "0", "0.0", None)]')
doc.add_paragraph()

heading2("Issue 4: translator.py _ENTITY_RULE 불완전")
body("파일: pipeline/translator.py:9-14", bold=True)
doc.add_paragraph()
body("문제")
body("회사·제품명 28개만 열거. guide v2 규칙 1이 요구하는 항목 누락:")
bullet("인물명: Jensen Huang, Sam Altman, Elon Musk, Sundar Pichai 등")
bullet("컨퍼런스명: NeurIPS, ICML, ICLR, CES, TechCrunch Disrupt 등")
bullet("지명: San Francisco, United States, California 등")
doc.add_paragraph()
body("영향")
body("Ollama 기반 프로덕션 translator가 인물명·컨퍼런스·지명을 한국어로 음차할 수 있음.")
doc.add_paragraph()
body("수정 내용")
body("_ENTITY_RULE을 [회사], [모델·제품], [인물], [컨퍼런스·행사] 4개 카테고리로 확장.")
doc.add_paragraph()

heading2("Issue 5: 베이스/파인튜닝 TPR 비교 비대칭 (설계 결정 필요)")
body("파일: eval/run_eval_finetuned.py, eval/run_eval_base.py", bold=True)
doc.add_paragraph()
body("문제")
bullet("파인튜닝 eval: translate() 내부에서 restore_entities() 후처리 적용 → 음차 자동 복원 → TPR 상승")
bullet("베이스 eval: restore_entities() 미적용 → 음차 그대로 측정 → TPR 낮음")
doc.add_paragraph()
body("영향")
body("파인튜닝 전/후 TPR 개선폭이 '모델+후처리' vs '모델만'의 비교가 됨 → 개선율 과대 계상.")
doc.add_paragraph()
body("권장 처리 (팀 결정 필요)")
bullet("옵션 A: 베이스 평가에도 restore_entities() 적용 (순수 모델 성능 공정 비교)")
bullet("옵션 B: 현행 유지 + '파인튜닝 모델은 후처리 포함' 명시 (실제 서비스 조건 반영)")
doc.add_paragraph()

# ══════════════════════════════════════════════════════
# 4. MINOR 이슈
# ══════════════════════════════════════════════════════
heading1("4. MINOR — 코드 품질")

heading2("Issue 6: bleu_comet.py 불필요한 중복 import")
body("파일: eval/metrics/bleu_comet.py:85", bold=True)
code_block("if model is None:\n    from comet import download_model, load_from_checkpoint  # 79행에 이미 동일 import\n    model = load_from_checkpoint(download_model(model_name))")
body("수정: 중복 import 제거 완료.")
doc.add_paragraph()

heading2("Issue 7: geval.py JSON 파싱 실패 시 재시도 없이 즉시 반환")
body("파일: eval/metrics/geval.py:214-221", bold=True)
code_block("except (json.JSONDecodeError, KeyError):\n    return {0점...}  # retries 루프 탈출, 재시도 없음")
body("파싱 실패 원인이 모델 응답 포맷 불안정(랜덤성)일 경우 재시도 시 성공 가능.")
body("수정: continue로 재시도 루프 유지, 마지막 시도에서만 0 반환으로 변경 완료.")
doc.add_paragraph()

# ══════════════════════════════════════════════════════
# 5. 수정 우선순위 요약
# ══════════════════════════════════════════════════════
heading1("5. 수정 우선순위 요약")

add_table(
    ["우선순위", "파일", "이슈", "예상 효과", "상태"],
    [
        ("1 (즉시)", "eval/metrics/geval.py", "Fluency 기준 guide v2로 교체", "G-Eval 점수 직접 상승", "✅ 완료"),
        ("2 (즉시)", "pipeline/translate_summarize.py", "Rule 3·4 guide v2로 통일", "프로덕션 일관성, 베이스 TPR 정확화", "✅ 완료"),
        ("3 (권장)", "run_eval_base.py + run_eval_finetuned.py", "TPR mean() 필터 수정", "보고 수치 정확화", "✅ 완료"),
        ("4 (권장)", "pipeline/translator.py", "_ENTITY_RULE 인물명·컨퍼런스 추가", "프로덕션 TPR 향상", "✅ 완료"),
        ("5 (낮음)", "eval/metrics/bleu_comet.py", "중복 import 제거", "코드 정리", "✅ 완료"),
        ("6 (낮음)", "eval/metrics/geval.py", "파싱 실패 재시도 복원", "G-Eval 안정성 향상", "✅ 완료"),
        ("7 (팀결정)", "eval/ 두 스크립트", "TPR 비교 비대칭 처리 방식 확정", "평가 공정성", "⏳ 미결"),
    ]
)

# ══════════════════════════════════════════════════════
# 6. 확인 필요
# ══════════════════════════════════════════════════════
heading1("6. 추가 확인 필요 항목")
bullet("pipeline/utils.py — preprocess_text() 동작 확인 (summarizer.py에서 사용)")
bullet("supabase_schema.sql — title 컬럼 코멘트 업데이트 완료 여부")
bullet("Issue 5 (TPR 비교 비대칭) — 팀 논의 후 처리 방식 결정")

# ── 저장 ──────────────────────────────────────────────
doc.save(OUTPUT)
print(f"저장 완료: {OUTPUT}")
