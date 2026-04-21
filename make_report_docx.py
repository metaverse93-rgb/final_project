"""삼선뉴스 파인튜닝 평가 개선 보고서 DOCX 생성"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# 기본 스타일
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10)

def set_font(run, bold=False, size=10, color=None):
    run.font.name = '맑은 고딕'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        set_font(run, bold=True, size=14, color=(0x1a, 0x56, 0xDB))
    elif level == 2:
        set_font(run, bold=True, size=12, color=(0x1E, 0x40, 0xAF))
    else:
        set_font(run, bold=True, size=10)
    return p

def shade_cell(cell, hex_color='DBEAFE'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            set_font(run, bold=True, size=9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade_cell(cell)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                set_font(run, size=9)
    return table

def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    set_font(r, size=10)

def note(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=9, color=(0x6B, 0x72, 0x80))

# ══════════════════════════════════════════
# 표지
# ══════════════════════════════════════════
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('삼선뉴스 파인튜닝 평가 개선 보고서')
set_font(r, bold=True, size=18, color=(0x1a, 0x56, 0xDB))

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('testset 200건 베이스라인 vs 파인튜닝 비교  |  TPR / G-Eval 개선 전략 적용')
set_font(r2, size=11, color=(0x6B, 0x72, 0x80))

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('작성일: 2026-04-17   |   담당: 이동우   |   모델: Qwen3.5-4B (QLoRA)')
set_font(r3, size=10, color=(0x6B, 0x72, 0x80))
doc.add_paragraph()

# ══════════════════════════════════════════
# 1. 평가 개요
# ══════════════════════════════════════════
heading(doc, '1. 평가 개요', 1)
p = doc.add_paragraph()
r = p.add_run(
    '삼선뉴스 프로젝트의 번역·요약 모델(Qwen3.5-4B QLoRA 파인튜닝)에 대해 '
    'testset 200건을 기준으로 베이스라인과 파인튜닝 모델의 성능을 비교하고, '
    'TPR 및 G-Eval 지표 개선을 위한 전략을 적용하여 결과를 정리합니다.'
)
set_font(r, size=10)

heading(doc, '1-1. 평가 지표 정의', 2)
add_table(doc,
    ['지표', '설명', '목표'],
    [
        ['BLEU', '번역 n-gram 정밀도 (기계번역 표준 지표)', '>= 17.0'],
        ['COMET', '신경망 기반 의미론적 번역 품질 (Unbabel wmt22-comet-da)', '베이스라인 대비 향상'],
        ['TPR', 'AI 신조어·고유명사 영문 보존율 (Nvidia, OpenAI 등 34개 용어 체크)', '>= 95%'],
        ['G-Eval', 'Claude Haiku 기반 4축 요약 품질 (일치성·유창성·일관성·관련성)', '>= 4.0'],
    ]
)
doc.add_paragraph()

# ══════════════════════════════════════════
# 2. 최종 평가 결과
# ══════════════════════════════════════════
heading(doc, '2. 최종 평가 결과 (testset 200건)', 1)
add_table(doc,
    ['지표', 'Base (200건)', 'Finetuned (200건)', '변화', '목표', '달성여부'],
    [
        ['BLEU',   '13.16', '23.96', '+10.80', '>= 17.0', '달성 (O)'],
        ['COMET',  '0.8721', '0.8741', '+0.0020', '베이스 대비 향상', '달성 (O)'],
        ['TPR',    '41.3%', '85.9%', '+44.6%p', '>= 95%', '미달 (-9.1%p)'],
        ['G-Eval', '3.83', '3.64', '-0.19', '>= 4.0', '미달'],
    ]
)
doc.add_paragraph()

heading(doc, '2-1. G-Eval 세부 지표 (Finetuned 200건)', 2)
add_table(doc,
    ['축', '점수', '목표', '해석'],
    [
        ['일치성 (Consistency)', '3.48', '-', '요약 내용이 원문과 사실적으로 일치하는 정도'],
        ['유창성 (Fluency)',     '3.85', '-', '문장의 자연스러움·문법적 완성도'],
        ['일관성 (Coherence)',   '4.01', '4.0 (달성)', '문장 간 논리적 흐름과 주제 일관성'],
        ['관련성 (Relevance)',   '3.22', '-', '원문 핵심 주제와의 연관성 (가장 낮음 — 개선 대상)'],
        ['G-Eval 평균',          '3.64', '>= 4.0', '4축 단순 평균'],
    ]
)
doc.add_paragraph()

# ══════════════════════════════════════════
# 3. 문제 원인 분석
# ══════════════════════════════════════════
heading(doc, '3. TPR·G-Eval 미달 원인 분석', 1)

heading(doc, '3-1. TPR 85.9% 미달 원인', 2)
bullet(doc, '학습 데이터(trainset_translate_1000_chat.jsonl)에 고유명사 음역 예시가 일부 혼입')
bullet(doc, '모델이 Nvidia→엔비디아, OpenAI→오픈에이아이 등 음역 패턴을 과적합')
bullet(doc, '기존 시스템 프롬프트에 "절대 음역 금지" 명시 없이 긍정 예시만 제공')

heading(doc, '3-2. G-Eval Relevance 3.22 미달 원인', 2)
bullet(doc, '요약 입력이 번역문(KO)만이었기 때문에 원문(EN) 핵심 주제가 희석됨')
bullet(doc, '파인튜닝 모델이 번역문 표현을 따라가며 원문 핵심 대신 번역 표면을 요약하는 경향 발생')
bullet(doc, '근거: G-Eval 논문(Liu et al., 2023 EMNLP) — Relevance는 source article과의 alignment로 측정')
doc.add_paragraph()

# ══════════════════════════════════════════
# 4. 개선 전략 및 적용 내용
# ══════════════════════════════════════════
heading(doc, '4. 개선 전략 및 적용 내용', 1)

heading(doc, '4-1. TPR 개선 전략', 2)
add_table(doc,
    ['전략', '적용 내용', '논문 근거', '상태'],
    [
        ['프롬프트 강화',
         'TRANSLATE_SYSTEM에 CRITICAL 키워드 추가\n절대 음역 금지 목록 28개 명시',
         'CTRL (Keskar et al., 2019)\n명시적 제어 신호가 규칙 준수율 향상',
         '완료'],
        ['후처리 복원',
         'restore_entities() 함수 신규 작성\n12개 음역 패턴 regex로 영문 복원\n(eval/metrics/term_preservation.py)',
         'Ugawa et al., 2018 EMNLP\nNMT Named Entity Preservation 후처리 표준',
         '완료'],
        ['학습 데이터 정제',
         'trainset에서 고유명사 음역 샘플 필터링\n수정 후 재파인튜닝',
         'Junczys-Dowmunt et al., 2018\n학습 데이터 노이즈 제거가 용어 일관성에 가장 큰 영향',
         '6주차 검토'],
    ]
)
doc.add_paragraph()

heading(doc, '4-2. G-Eval 개선 전략', 2)
add_table(doc,
    ['전략', '적용 내용', '논문 근거', '상태'],
    [
        ['원문 기반 요약',
         'summarize() 입력을 EN원문+KO번역문\n동시 제공으로 변경\nRelevance 직접 개선',
         'G-Eval (Liu et al., 2023 EMNLP)\nRelevance는 source alignment가 핵심',
         '완료'],
        ['요약 프롬프트 재설계',
         'SUMMARIZE_SYSTEM에\n"원문 핵심 주제·사실에 충실" 명시\nRelevance 최우선 명시',
         '동일',
         '완료'],
        ['요약 데이터 품질 보강',
         'GT summary가 원문 핵심 문장과\nhigh overlap인 샘플 선별 후 재파인튜닝',
         'PEGASUS (Zhang et al., 2020 ICML)\nGT-원문 overlap이 Relevance 학습에 유리',
         '6주차 검토'],
    ]
)
doc.add_paragraph()

# ══════════════════════════════════════════
# 5. 수정 파일 목록
# ══════════════════════════════════════════
heading(doc, '5. 수정된 파일 상세', 1)
add_table(doc,
    ['파일 경로', '변경 내용'],
    [
        ['eval/metrics/term_preservation.py',
         'restore_entities() 추가\n12개 음역 패턴 regex 후처리 (엔비디아->Nvidia 등)'],
        ['eval/run_eval_finetuned.py',
         '(1) TRANSLATE_SYSTEM: CRITICAL 고유명사 절대 보존 규칙 강화\n'
         '(2) SUMMARIZE_SYSTEM: 원문 핵심 주제 기반 요약 명시\n'
         '(3) translate(): restore_entities() 후처리 연결\n'
         '(4) summarize(en_text, ko_text): 원문+번역문 동시 입력으로 변경\n'
         '(5) --output 옵션 추가: 별도 파일로 검증 재실행 가능'],
        ['pipeline/translator.py',
         '프로덕션 번역 프롬프트 강화\n고유명사 28개 영문 보존 규칙 명시'],
        ['pipeline/summarizer.py',
         'summarize(en_text, ko_text) 시그니처 변경\n원문 기반 요약으로 전환'],
        ['backend/save_articles.py',
         'title_ko 필드 추가\ntranslate_summarize 출력의 한국어 제목을 DB에 저장'],
        ['supabase_schema.sql',
         'articles 테이블에 title_ko TEXT 컬럼 추가'],
    ]
)
doc.add_paragraph()

# ══════════════════════════════════════════
# 6. 샘플 검증 결과
# ══════════════════════════════════════════
heading(doc, '6. 개선 적용 후 샘플 검증 결과 (5건)', 1)
add_table(doc,
    ['지표', '기존 v1 (200건)', '개선 v2 (5건 샘플)', '변화'],
    [
        ['BLEU',   '23.96', '18.31', '샘플 편차 있음 (소량 기준)'],
        ['COMET',  '0.8741', '0.9087', '+0.034 향상'],
        ['TPR',    '85.9%', '90.0%', '+4.1%p 향상'],
        ['G-Eval', '3.64', '3.95', '+0.31 향상 (목표 4.0 근접)'],
    ]
)
note(doc, '※ 5건 샘플 기준이므로 통계적 신뢰도는 낮습니다. --skip-geval 옵션으로 전체 200건 재평가 진행 예정.')
doc.add_paragraph()

# ══════════════════════════════════════════
# 7. title_ko DB 저장
# ══════════════════════════════════════════
heading(doc, '7. title_ko (한국어 제목) DB 저장 현황', 1)
p = doc.add_paragraph()
r = p.add_run(
    'pipeline/translate_summarize.py 에서 title_ko 필드가 이미 생성되고 있었으나 '
    'backend/save_articles.py 및 supabase_schema.sql 에 미반영 상태였습니다. '
    '이번 세션에서 수정 완료했습니다.'
)
set_font(r, size=10)
add_table(doc,
    ['파일', '수정 전', '수정 후'],
    [
        ['supabase_schema.sql', 'title TEXT만 존재', 'title_ko TEXT 컬럼 추가'],
        ['backend/save_articles.py', 'title_ko 저장 없음', 'title_ko: a.get("title_ko") 저장'],
    ]
)
p2 = doc.add_paragraph()
r2 = p2.add_run('※ 적용 방법: Supabase SQL Editor에서 아래 쿼리 실행 필요')
set_font(r2, size=9, color=(0xDC, 0x26, 0x26))
p3 = doc.add_paragraph()
r3 = p3.add_run('ALTER TABLE articles ADD COLUMN IF NOT EXISTS title_ko TEXT;')
set_font(r3, size=9)
p3.paragraph_format.left_indent = Cm(1)
doc.add_paragraph()

# ══════════════════════════════════════════
# 8. 향후 계획
# ══════════════════════════════════════════
heading(doc, '8. 향후 계획 (6주차)', 1)
add_table(doc,
    ['날짜', '작업', '목표'],
    [
        ['04/20 (월)', '--skip-geval 전체 200건 재평가 (개선 프롬프트 적용)', 'TPR >= 95% 확인'],
        ['04/21 (화)', '신뢰도 분류 수동 검증 100건', 'FACT/RUMOR/UNVERIFIED 정확도 측정'],
        ['04/22 (수)', '맥락 연결 UI 요약·번역 출력 연동', '관련 기사 3개 출력 확인'],
        ['04/23 (목)', '평가 보고서 초안 작성 (수치 포함)', '보고서 완성'],
        ['검토 중', '학습 데이터 정제 + 재파인튜닝', 'TPR·G-Eval 근본 해결'],
    ]
)
doc.add_paragraph()

out = r'C:\Users\이동우\Desktop\삼선뉴스_파인튜닝평가_개선보고서_0417.docx'
doc.save(out)
print(f'저장 완료: {out}')
